from __future__ import annotations

import io
import os
import zipfile
from typing import Any, Dict, List

from docx import Document as DocxDocument
from PIL import Image
from llama_index.core import Document

from states.loaders.utils import analyze_image_with_lvm, run_ocr_on_pil, save_pil_image


def load_docx(path: str, filename: str, artifacts: Dict[str, List[Any]] | None = None) -> List[Document]:
    try:
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        full_text = "\n".join(paragraphs).strip() or "[EMPTY DOCX]"
        docs: List[Document] = [Document(text=full_text, metadata={"filename": filename, "type": "docx"})]

        try:
            with zipfile.ZipFile(path) as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/") and not name.endswith("/"):
                        img_bytes = zf.read(name)
                        pil_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                        img_path = save_pil_image(pil_img, f"{filename}_docx_{os.path.basename(name)}")
                        caption, insights = analyze_image_with_lvm(pil_img)
                        ocr_text = run_ocr_on_pil(pil_img)
                        if artifacts is not None:
                            artifacts["extracted_images"].append(img_path)
                            artifacts["image_descriptions"].append(caption)
                            artifacts["image_insights"].append(insights)
                        docs.append(
                            Document(
                                text=(
                                    "[DOCX IMAGE]\n"
                                    f"Filename:{filename}\nImageName:{name}\n"
                                    f"Caption:{caption}\nInsights:{insights}\nOCR:{ocr_text}"
                                ),
                                metadata={"filename": filename, "type": "docx-image", "image_path": img_path},
                            )
                        )
        except Exception:
            pass

        return docs
    except Exception as e:
        print(f"Failed to load DOCX {filename}: {e}")
        return []

